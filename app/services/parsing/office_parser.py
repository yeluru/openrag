"""Office Open XML: .docx (Word) and .xlsx (Excel) → :class:`ParsedDocument`."""

from __future__ import annotations

import logging
from pathlib import Path

from docx import Document as DocxDocument
from openpyxl import load_workbook

from app.services.parsing.models import ParsedDocument, ParsedPage
from app.services.parsing.registry import register_parser

logger = logging.getLogger(__name__)


def _pkg_version(dist_name: str) -> str:
    try:
        from importlib.metadata import version

        return version(dist_name)
    except Exception:
        return "0"


_DOCX_CANONICAL = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
_XLSX_CANONICAL = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"


class DocxParser:
    name = "python-docx"
    version = _pkg_version("python-docx")

    def parse(self, file_path: str) -> ParsedDocument:
        path = Path(file_path)
        doc = DocxDocument(str(path))
        parts: list[str] = []
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if t:
                parts.append(t)
        for table in doc.tables:
            for row in table.rows:
                cells = [(c.text or "").strip() for c in row.cells]
                line = " | ".join(x for x in cells if x)
                if line:
                    parts.append(line)
        text = "\n\n".join(parts)
        limitations: list[str] = []
        if not text.strip():
            limitations.append("No extractable text in this Word document.")

        meta: dict = {"title": None, "author": None}
        try:
            cp = doc.core_properties
            if cp:
                meta["title"] = cp.title
                meta["author"] = cp.author
        except Exception as e:
            logger.debug("docx core_properties: %s", e)

        page = ParsedPage(page_number=1, text=text if text.strip() else "(empty)", blocks=[])
        return ParsedDocument(
            pages=[page],
            file_path=str(path),
            parser_name=self.name,
            parser_version=self.version,
            limitations=limitations,
            raw_metadata=meta,
        )


class XlsxParser:
    name = "openpyxl"
    version = _pkg_version("openpyxl")

    def parse(self, file_path: str) -> ParsedDocument:
        path = Path(file_path)
        wb = load_workbook(str(path), read_only=True, data_only=True)
        pages: list[ParsedPage] = []
        limitations: list[str] = []
        try:
            for i, sheetname in enumerate(wb.sheetnames, start=1):
                ws = wb[sheetname]
                rows_text: list[str] = []
                for row in ws.iter_rows(values_only=True):
                    cells = ["" if c is None else str(c).strip() for c in row]
                    if not any(cells):
                        continue
                    rows_text.append("\t".join(cells))
                body = "\n".join(rows_text)
                header = f"Sheet: {sheetname}\n\n" if sheetname else ""
                sheet_text = (header + body).strip() or f"(empty sheet: {sheetname})"
                pages.append(ParsedPage(page_number=i, text=sheet_text, blocks=[]))
        finally:
            wb.close()

        if not pages:
            limitations.append("Workbook has no sheets.")
            pages = [ParsedPage(page_number=1, text="(empty workbook)", blocks=[])]

        return ParsedDocument(
            pages=pages,
            file_path=str(path),
            parser_name=self.name,
            parser_version=self.version,
            limitations=limitations,
            raw_metadata={"sheet_count": len(pages)},
        )


def _register() -> None:
    register_parser(
        canonical_mime=_DOCX_CANONICAL,
        mime_types=frozenset({_DOCX_CANONICAL}),
        extensions=frozenset({".docx"}),
        factory=lambda: DocxParser(),
    )
    register_parser(
        canonical_mime=_XLSX_CANONICAL,
        mime_types=frozenset({_XLSX_CANONICAL}),
        extensions=frozenset({".xlsx"}),
        factory=lambda: XlsxParser(),
    )


_register()
